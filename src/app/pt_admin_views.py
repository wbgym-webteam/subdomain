#subdomain\src\app\pt_admin_views.py
from flask import (
    Blueprint,
    render_template,
    request,
    redirect,
    url_for,
    session,
    send_from_directory,
    send_file,
    Response,
    current_app,
    jsonify
)

from sqlalchemy import text

import json
import os
import zipfile
import html
import traceback
from .pt_filehandler import FileHandlerPT

from .models import PTStudent, PTPresentation, PTSelection
from . import db
from .pt_logincode_export import export_logincodes as export_logincodesPT
from .pt_selection_export import SelectionExporter as SelectionExporterPT
from .pt_selection_engine import run_pt_selection_generator

# secure used, because we can't store Personal Data, so secure version of the filehandler is used
from .pt_filehandler import load_names_map

# secure used, because we can't store Personal Data, so secure version of the filehandler is used
from .pt_filehandler_secure import FileHandlerPTSecure, load_names_map

import io

pt_admin_views = Blueprint("pt_admin_views", __name__, static_folder="static")

# ------------------------------------------------------------------
# Decorator
from functools import wraps


def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get("admin_logged_in"):
            # Check if this is a fetch/AJAX request
            # Fetch requests typically don't have X-Requested-With header,
            # but we can check for common patterns
            is_fetch = (
                request.headers.get('X-Requested-With') == 'XMLHttpRequest' or
                request.is_json or
                'application/json' in request.headers.get('Accept', '') or
                request.headers.get('Content-Type', '').startswith('application/json')
            )

            if is_fetch:
                # Return 401 Unauthorized for AJAX/fetch requests
                return jsonify({'error': 'Unauthorized', 'redirect': '/admin_login'}), 401

            # Preserve the original URL for redirect after login (for normal page requests)
            return redirect(f"/admin_login?next={request.url}")
        return f(*args, **kwargs)

    return decorated_function


# ------------------------------------------------------------------
# Routing for PT admin views

@pt_admin_views.route("/pt/panel", methods=["GET", "POST"])
@admin_required
def pt_Panel():
    current_dir = os.path.dirname(os.path.abspath(__file__))
    module_status_path = os.path.join(current_dir, "data", "module_status.json")
    with open(module_status_path, "r") as f:
        module_status = json.load(f)
        ms = module_status["modules"]["PT"]
    return render_template("admin/pt_panel.html", status=ms)


@pt_admin_views.route("/pt/upload_file", methods=["POST"])
@admin_required
def pt_upload_file():
    try:
        if "file" not in request.files:
            session['upload_error'] = "No file provided"
            return redirect("/admin/pt/panel")
        file = request.files["file"]
        if file.filename == "":
            session['upload_error'] = "Empty filename"
            return redirect("/admin/pt/panel")

        # Use absolute path from app root
        current_dir = os.path.dirname(os.path.abspath(__file__))
        upload_dir = os.path.join(current_dir, 'data', 'pt', 'uploads')
        os.makedirs(upload_dir, exist_ok=True)

        # Check if this is a names file upload
        file_type = request.form.get("file_type", "data")

        if file_type == "names":
            # Save as names file
            file_path = os.path.join(upload_dir, "names_workbook.xlsx")
            file.save(file_path)
            print(f"Names file uploaded successfully to {file_path}")
            session['upload_success'] = "Names file uploaded successfully"
        else:
            # Save as regular data file
            file_path = os.path.join(upload_dir, "workbook.xlsx")
            file.save(file_path)
            print(f"File saved to {file_path}")

            # Process the uploaded file
            FileHandlerPTSecure()
            print("Data file uploaded and processed successfully")
            session['upload_success'] = "Data file uploaded and processed successfully"

        return redirect("/admin/pt/panel")

    except Exception as e:
        print(f"Error in pt_upload_file: {e}")
        print(traceback.format_exc())
        session['upload_error'] = f"Upload failed: {str(e)}"
        return redirect("/admin/pt/panel")


@pt_admin_views.route("/pt/module_status", methods=["POST"])
@admin_required
def pt_module_status():
    current_dir = os.path.dirname(os.path.abspath(__file__))
    module_status_path = os.path.join(current_dir, "data", "module_status.json")

    with open(module_status_path, "r") as f:
        data = json.load(f)

    current_status = data["modules"]["PT"]

    if current_status == "active":
        data["modules"]["PT"] = "inactive"
    else:
        data["modules"]["PT"] = "active"

    with open(module_status_path, "w") as f:
        json.dump(data, f, indent=4)

    return redirect("/admin/pt/panel")

# Version that requires the Names to be Saved in the DB
@pt_admin_views.route("/pt/export_logincodes", methods=["POST"])
@admin_required
def pt_export_logincodes_route():
    if request.method == "POST":
        export_logincodesPT()
        return redirect("/admin/pt/panel")  # Use absolute path

# Version that loads names from uploaded file in RAM ONLY and does not require Names to be saved in DB 
@pt_admin_views.route("/pt/export_logincodes_secure", methods=["POST"])
@admin_required
def pt_export_logincodes_secure():
    if "file" not in request.files:
        return redirect("/admin/pt/panel")
    
    file = request.files["file"]
    
    try:
        import io
        import zipfile
        from docx import Document
        from docx.shared import Pt
        from .pt_filehandler import load_names_map
        
        # 1. Load Names into MEMORY
        names_map = load_names_map(file)
        
        # 2. Prepare ZIP Buffer
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            
            # --- STRATEGY 1: Grades 5, 6, 11, 12 (Export as whole grades) ---
            special_grades = [5, 6, 11, 12]
            
            for grade in special_grades:
                # Get IDs only
                students = db.session.execute(text(
                    f"SELECT id, logincode FROM pt_students WHERE grade = {grade}"
                )).all()
                
                if not students:
                    continue
                    
                # Merge Names & Sort
                student_data = []
                for s in students:
                    n = names_map.get(s.id, {'first': 'Unknown', 'last': 'Unknown'})
                    student_data.append({
                        'first': n['first'], 
                        'last': n['last'], 
                        'code': s.logincode
                    })
                
                # Sort by Last Name
                student_data.sort(key=lambda x: (x['last'], x['first']))
                
                # Generate DOCX in RAM
                doc = Document()
                doc.add_heading(f"PT Login Codes {grade}", 0)
                
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                
                # Headers
                headers = ["First Name", "Last Name", "Login Code"]
                for i, h in enumerate(headers):
                    cell = table.cell(0, i)
                    cell.text = h
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(14)
                
                # Rows
                for s in student_data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = s['first']
                    row_cells[1].text = s['last']
                    row_cells[2].text = s['code']
                    
                    # Apply spacing style to row
                    for cell in row_cells:
                        for p in cell.paragraphs:
                            p.paragraph_format.space_after = Pt(14)
                            p.paragraph_format.space_before = Pt(14)
                
                # Save to ZIP
                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                zip_file.writestr(f"PT_Logincodes_{grade}.docx", docx_buffer.getvalue())

            # --- STRATEGY 2: All other grades (Split by Selector) ---
            # Find which grades/selectors exist (excluding the special ones)
            combos = db.session.execute(text(
                f"SELECT DISTINCT grade, grade_selector FROM pt_students WHERE grade NOT IN (5, 6, 11, 12) ORDER BY grade, grade_selector"
            )).all()
            
            for grade, selector in combos:
                # Get IDs
                students = db.session.execute(text(
                    f"SELECT id, logincode FROM pt_students WHERE grade = {grade} AND grade_selector = {selector}"
                )).all()
                
                if not students:
                    continue

                # Merge Names & Sort
                student_data = []
                for s in students:
                    n = names_map.get(s.id, {'first': 'Unknown', 'last': 'Unknown'})
                    student_data.append({
                        'first': n['first'], 
                        'last': n['last'], 
                        'code': s.logincode
                    })
                
                student_data.sort(key=lambda x: (x['last'], x['first']))
                
                # Generate DOCX
                doc = Document()
                doc.add_heading(f"PT Login Codes {grade}/{selector}", 0)
                
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                
                # Headers
                headers = ["First Name", "Last Name", "Login Code"]
                for i, h in enumerate(headers):
                    cell = table.cell(0, i)
                    cell.text = h
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(14)
                
                # Rows
                for s in student_data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = s['first']
                    row_cells[1].text = s['last']
                    row_cells[2].text = s['code']
                    
                    for cell in row_cells:
                        for p in cell.paragraphs:
                            p.paragraph_format.space_after = Pt(14)
                            p.paragraph_format.space_before = Pt(14)

                # Save to ZIP
                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                zip_file.writestr(f"PT_Logincodes_{grade}_{selector}.docx", docx_buffer.getvalue())

        # 3. Finalize and Send
        zip_buffer.seek(0)
        return send_file(
            zip_buffer,
            as_attachment=True,
            download_name="PT_Logincodes_Secure.zip",
            mimetype="application/zip"
        )

    except Exception as e:
        print(f"Error in secure logincode export: {e}")
        import traceback
        traceback.print_exc()
        return redirect("/admin/pt/panel")

# Version that requires the Names to be Saved in the DB
@pt_admin_views.route("/pt/download_logincodes", methods=["GET"])
@admin_required
def pt_download_logincodes():  # Ensure this function is used for the pt route
    current_dir = os.path.dirname(os.path.abspath(__file__))
    download_dir = os.path.join(current_dir, 'data', 'pt', 'downloads')
    return send_from_directory(download_dir, "PT_Logincodes.zip", as_attachment=True)

# Version that requires the Names to be Saved in the DB
@pt_admin_views.route("/pt/export_wishes", methods=["POST"])
@admin_required
def pt_export_wishes():
    try:
        print("Starting PT wishes export...")
        result_file_path = SelectionExporterPT(db)
        if result_file_path and os.path.exists(result_file_path):
            print(f"Export successful, file created at: {result_file_path}")
            return redirect("/admin/pt/panel")
        else:
            print("Export failed - no file created")
            return redirect("/admin/pt/panel")
    except Exception as e:
        print(f"Error exporting wishes: {e}")
        import traceback
        traceback.print_exc()
        return redirect("/admin/pt/panel")  # Use absolute path

# Version that loads names from uploaded file in RAM ONLY and does not require Names to be saved in DB 
@pt_admin_views.route("/pt/export_wishes_secure", methods=["POST"])
@admin_required
def pt_export_wishes_secure():
    if "file" not in request.files:
        return redirect("/admin/pt/panel")
    
    file = request.files["file"]
    
    try:
        import io
        from openpyxl import Workbook
        from .pt_filehandler import load_names_map
        
        # 1. Load Names into MEMORY
        names_map = load_names_map(file)
        
        # 2. Query DB for Wishes (IDs only)
        # We join pt_selections, pt_students, and pt_presentations
        wishes_query = db.session.execute(text("""
            SELECT 
                s.id AS student_id, 
                s.grade, 
                s.grade_selector,
                s.logincode,
                p.id AS course_id, 
                p.title AS course_title,
                sel.ranking
            FROM pt_selections sel
            JOIN pt_students s ON sel.student_id = s.id
            JOIN pt_presentations p ON sel.presentation_id = p.id
            ORDER BY s.grade, s.grade_selector, s.id, sel.ranking
        """)).mappings().all()
        
        # 3. Build Excel in RAM
        wb = Workbook()
        ws = wb.active
        ws.title = "Student Wishes"
        
        # Headers
        ws.append([
            "Student ID", "Last Name", "First Name", "Grade", "Class", "Login Code",
            "Rank", "Course ID", "Course Title"
        ])
        
        # 4. Fill Data
        for row in wishes_query:
            s_id = row['student_id']
            
            # Lookup Name in RAM
            name_entry = names_map.get(s_id, {'first': 'Unknown', 'last': 'Unknown'})
            
            ws.append([
                s_id,
                name_entry['last'],
                name_entry['first'],
                row['grade'],
                row['grade_selector'],
                row['logincode'],
                row['ranking'],
                row['course_id'],
                row['course_title']
            ])
            
        # 5. Save to RAM Buffer
        excel_buffer = io.BytesIO()
        wb.save(excel_buffer)
        excel_buffer.seek(0)
        
        return send_file(
            excel_buffer,
            as_attachment=True,
            download_name="PT_Kurs_Wuensche_Secure.xlsx",
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    except Exception as e:
        print(f"Error in secure wishes export: {e}")
        import traceback
        traceback.print_exc()
        return redirect("/admin/pt/panel")

# Version that requires the Names to be Saved in the DB
@pt_admin_views.route("/pt/download_wishes", methods=["GET"])
@admin_required
def pt_download_wishes():
    try:
        current_dir = os.path.dirname(os.path.abspath(__file__))
        download_dir = os.path.join(current_dir, 'data', 'pt', 'downloads')
        file_name = "Kurs_Wuensche.xlsx"
        file_path = os.path.join(download_dir, file_name)
        print(f"Looking for file at: {file_path}")
        print(f"File exists: {os.path.exists(file_path)}")
        if not os.path.exists(file_path):
            print("File not found, creating export first...")
            result_file_path = SelectionExporterPT(db)
            if not result_file_path or not os.path.exists(result_file_path):
                print("Could not create export file")
                return redirect("/admin/pt/panel")
        return send_from_directory(
            download_dir,
            file_name,
            as_attachment=True,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    except Exception as e:
        print(f"Error downloading wishes: {e}")
        import traceback
        traceback.print_exc()
        return redirect("/admin/pt/panel")  # Use absolute path

@pt_admin_views.route("/pt/run_selection", methods=["POST"])
@admin_required
def pt_run_selection():
    """
    This route now returns a streaming response of *just*
    HTML log fragments, not a full page.
    """
    app_ctx = current_app._get_current_object()

    def stream_progress(app):
        # 4. Wrap the *entire* generator in the app context.
        with app.app_context():
            try:
                # 2. Run the generator and stream its output
                engine_generator = run_pt_selection_generator()
                for log_line in engine_generator:
                    # Escape the log line to prevent HTML injection
                    safe_line = html.escape(log_line)
                    
                    # Add some color coding for readability
                    if "ERROR" in log_line or "FATAL" in log_line or "WARNING" in log_line:
                        yield f'<p class="error">{safe_line}</p>'
                    elif "Iteration" in log_line:
                        yield f'<p class="iteration">{safe_line}</p>'
                    elif "--- Assignment Report ---" in log_line or "- " in log_line:
                        yield f'<p class="report">{safe_line}</p>'
                    elif "--- DONE ---" in log_line or "Successfully saved" in log_line:
                        yield f'<p class="success">{safe_line}</p>'
                    else:
                        yield f'<p>{safe_line}</p>'
                    
                    # Flush buffer
                    yield " " 
                
                yield '<p class="done">Engine run finished.</p>'

            except Exception as e:
                safe_error = html.escape(str(e))
                safe_traceback = html.escape(traceback.format_exc())
                yield f'<p class="error">A critical error occurred in the view: {safe_error}</p>'
                yield f'<p class="error">{safe_traceback}</p>'

    # Return the streaming response
    return Response(stream_progress(app_ctx), mimetype='text/html')

@pt_admin_views.route("/pt/view_assignments", methods=["GET"])
@admin_required
def pt_view_assignments():
    try:
        # Get all assignments with student and presentation details
        assignments_query = db.session.execute(
            text("""
                SELECT 
                    s.first_name, s.last_name, s.grade,
                    p.title, p.presenter, p.slot, p.column, p.room
                FROM pt_assignments a
                JOIN pt_students s ON a.student_id = s.id
                JOIN pt_presentations p ON a.presentation_id = p.id
                ORDER BY s.last_name, s.first_name, p.slot
            """)
        ).all()
        
        # Get assignment report from session if available
        report = session.get('assignment_report', {})
        
        return render_template(
            "admin/pt_assignments.html", 
            assignments=assignments_query,
            report=report
        )
    except Exception as e:
        print(f"Error viewing assignments: {e}")
        return redirect("/admin/pt/panel")

# Version that requires the Names to be Saved in the DB
@pt_admin_views.route("/pt/export_assignments", methods=["GET"])
@admin_required
def pt_export_assignments():
    try:
        import pandas as pd
        
        # Get all assignments with details
        assignments_query = db.session.execute(
            text("""
                SELECT 
                    s.first_name, s.last_name, s.grade, s.logincode,
                    p.title, p.presenter, p.teacher, p.description, 
                    p.slot, p.column, p.room, p.max_students
                FROM pt_assignments a
                JOIN pt_students s ON a.student_id = s.id
                JOIN pt_presentations p ON a.presentation_id = p.id
                ORDER BY s.last_name, s.first_name, p.slot
            """)
        ).all()
        
        # Convert to DataFrame
        df = pd.DataFrame(assignments_query, columns=[
            'First Name', 'Last Name', 'Grade', 'Login Code',
            'Course Title', 'Presenter', 'Teacher', 'Description',
            'Slot', 'Column', 'Room', 'Max Students'
        ])
        
        # Create download directory with absolute path
        current_dir = os.path.dirname(os.path.abspath(__file__))
        download_dir = os.path.join(current_dir, 'data', 'pt', 'downloads')
        os.makedirs(download_dir, exist_ok=True)
        
        # Save to Excel
        file_path = os.path.join(download_dir, 'PT_Assignments.xlsx')
        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='All Assignments', index=False)
            
            # Create separate sheets for each slot
            for slot in [1, 2, 3]:
                slot_df = df[df['Slot'] == slot]
                slot_df.to_excel(writer, sheet_name=f'Slot {slot}', index=False)
        
        return send_from_directory(
            download_dir,
            'PT_Assignments.xlsx',
            as_attachment=True,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        
    except Exception as e:
        print(f"Error exporting assignments: {e}")
        import traceback
        traceback.print_exc()
        return redirect("/admin/pt/panel")

# Version that requires the Names to be Saved in the DB
@pt_admin_views.route("/pt/export_assignments_pdf", methods=["GET"])
@admin_required
def pt_export_assignments_pdf():
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        current_dir = os.path.dirname(os.path.abspath(__file__))
        download_dir = os.path.join(current_dir, 'data', 'pt', 'downloads')
        os.makedirs(download_dir, exist_ok=True)
        
        zip_filename = 'PT_Schüler_Stundenpläne.zip'
        zip_path = os.path.join(download_dir, zip_filename)

        print(f"Cleaning old files from {download_dir}...")
        if os.path.exists(zip_path):
            os.remove(zip_path)
        for file in os.listdir(download_dir):
            if file.startswith('PT_Stundenpläne_') and file.endswith('.pdf'):
                os.remove(os.path.join(download_dir, file))
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.black
        )
        student_title_style = ParagraphStyle(
            'StudentTitle',
            parent=styles['Heading2'],
            fontSize=18,
            spaceAfter=20,
            alignment=TA_CENTER,
            textColor=colors.black
        )
        info_style = ParagraphStyle(
            'InfoStyle',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=10,
            alignment=TA_LEFT
        )
        
        classes_query = db.session.execute(
            text("""
                SELECT DISTINCT s.grade, s.grade_selector
                FROM pt_students s
                JOIN pt_assignments a ON s.id = a.student_id
                ORDER BY s.grade, s.grade_selector
            """)
        ).all()
        
        if not classes_query:
            print("No students with assignments found")
            return redirect("/admin/pt/panel") 
        
        print(f"Found {len(classes_query)} classes with assignments.")
        
        for grade, grade_selector in classes_query:
            
            pdf_filename = f"PT_Stundenpläne_{grade}_{grade_selector}.pdf"
            pdf_path = os.path.join(download_dir, pdf_filename)
            doc = SimpleDocTemplate(pdf_path, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
            story = []

            students_query = db.session.execute(
                text("""
                    SELECT DISTINCT s.id, s.first_name, s.last_name, s.grade, s.grade_selector
                    FROM pt_students s
                    JOIN pt_assignments a ON s.id = a.student_id
                    WHERE s.grade = :grade AND s.grade_selector = :grade_selector
                    ORDER BY s.last_name, s.first_name
                """),
                {"grade": grade, "grade_selector": grade_selector}
            ).all()

            if not students_query:
                continue

            story.append(Paragraph(f"PT-Kurszuordnungen - Klasse {grade}/{grade_selector}", title_style))
            story.append(Paragraph(f"{len(students_query)} Schüler", styles['Heading2']))
            story.append(PageBreak())
            
            for i, student in enumerate(students_query):
                student_id, first_name, last_name, s_grade, s_grade_selector = student
                
                # Add student header
                story.append(Paragraph(f"{last_name}, {first_name}", student_title_style))
                story.append(Paragraph(f"Klasse: {s_grade} | Klassenbezeichnung: {s_grade_selector}", info_style))
                story.append(Spacer(1, 0.5*cm))
                
                # Get student's assignments
                assignments_query = db.session.execute(
                    text("""
                        SELECT p.title, p.presenter, p.teacher, p.slot, p.room, p.description
                        FROM pt_assignments a
                        JOIN pt_presentations p ON a.presentation_id = p.id
                        WHERE a.student_id = :student_id
                        ORDER BY p.slot
                    """),
                    {"student_id": student_id}
                ).all()
                
                if assignments_query:
                    table_data = [['Block', 'Kurstitel', 'Referent', 'Lehrer', 'Raum']]
                    for assignment in assignments_query:
                        title, presenter, teacher, slot, room, description = assignment
                        table_data.append([
                            f"Block {slot}",
                            title or "N/A",
                            presenter or "N/A",
                            teacher or "N/A",
                            room or "N/A"
                        ])
                    
                    table = Table(table_data, colWidths=[2*cm, 6*cm, 3*cm, 3*cm, 2*cm])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 10),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 1*cm))
                    
                    has_descriptions = any(a[5] for a in assignments_query if a[5])
                    if has_descriptions:
                        story.append(Paragraph("Kursbeschreibungen:", styles['Heading3']))
                        story.append(Spacer(1, 0.3*cm))
                        for assignment in assignments_query:
                            title, presenter, teacher, slot, room, description = assignment
                            if description and description.strip():
                                desc_text = f"<b>Block {slot} - {title}:</b> {description}"
                                story.append(Paragraph(desc_text, info_style))
                                story.append(Spacer(1, 0.2*cm))
                else:
                    story.append(Paragraph("Keine Kurszuordnungen für diesen Schüler gefunden.", info_style))
                
                if i < len(students_query) - 1:
                    story.append(PageBreak())
            
         
            doc.build(story)
            print(f"Created PDF for class {grade}/{grade_selector} at {pdf_path}")


        print(f"Zipping PDFs into {zip_path}...")
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file in os.listdir(download_dir):
                if file.startswith('PT_Stundenpläne_') and file.endswith('.pdf'):
                    file_full_path = os.path.join(download_dir, file)
                    zipf.write(file_full_path, file)
                    os.remove(file_full_path) 

        print("Zipping complete.")

        return send_from_directory(
            download_dir,
            zip_filename,
            as_attachment=True,
            mimetype="application/zip"
        )
        
    except ImportError as e:
        print(f"Missing reportlab library: {e}")
        print("Please install reportlab: pip install reportlab")
        return redirect("/admin/pt/panel")
    except Exception as e:
        print(f"Error exporting PDF: {e}")
        import traceback
        traceback.print_exc()
        return redirect("/admin/pt/panel")

# Version that loads names from uploaded file in RAM ONLY and does not require Names to be saved in DB 
@pt_admin_views.route("/pt/export_assignments_pdf_secure", methods=["POST"])
@admin_required
def pt_export_assignments_pdf_secure():
    print("=" * 50)
    print("FUNCTION CALLED: pt_export_assignments_pdf_secure")
    print(f"Request method: {request.method}")
    print(f"Request files: {request.files}")
    print(f"Request form: {request.form}")
    print("=" * 50)

    if "file" not in request.files:
        error_msg = "No file uploaded for PDF export"
        print(f"ERROR: {error_msg}")
        session['upload_error'] = error_msg
        return redirect("/admin/pt/panel")

    file = request.files["file"]
    print(f"File received: {file.filename}")

    if file.filename == "":
        error_msg = "Empty filename for PDF export"
        print(f"ERROR: {error_msg}")
        session['upload_error'] = error_msg
        return redirect("/admin/pt/panel")

    try:
        # Check for required libraries first
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER
            print("ReportLab imported successfully")
        except ImportError as ie:
            error_msg = f"Missing required library: {str(ie)}. Please install reportlab: uv pip install reportlab"
            print(f"ERROR: {error_msg}")
            session['upload_error'] = error_msg
            return redirect("/admin/pt/panel")

        # 1. Load Names into MEMORY
        from .pt_filehandler import load_names_map
        names_map = load_names_map(file)
        print(f"Loaded {len(names_map)} names into memory.")

        if not names_map:
            session['upload_error'] = "Failed to load names from Excel file. Please check file format (ID, Last Name, First Name in columns A, B, C)"
            return redirect("/admin/pt/panel")

        # 3. Master Buffer for the ZIP file
        zip_buffer = io.BytesIO()

        # 4. Get Classes
        classes_query = db.session.execute(
            text("""
                SELECT DISTINCT s.grade, s.grade_selector
                FROM pt_students s
                JOIN pt_assignments a ON s.id = a.student_id
                ORDER BY s.grade, s.grade_selector
            """)
        ).all()

        print(f"Found {len(classes_query)} classes with assignments")

        if not classes_query:
            session['upload_error'] = "No assignments found in database. Please run the selection engine first."
            return redirect("/admin/pt/panel")

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24, spaceAfter=30, alignment=TA_CENTER)
        student_title_style = ParagraphStyle('StudentTitle', parent=styles['Heading2'], fontSize=18, spaceAfter=20, alignment=TA_CENTER)
        info_style = ParagraphStyle('InfoStyle', parent=styles['Normal'], fontSize=12, spaceAfter=10)

        # Open the ZIP file in memory
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            
            for grade, grade_selector in classes_query:
                # Prepare a buffer for THIS specific class PDF
                class_pdf_buffer = io.BytesIO()
                
                doc = SimpleDocTemplate(class_pdf_buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
                story = []

                # Get students (IDs only)
                students_query = db.session.execute(
                    text("""
                        SELECT DISTINCT s.id, s.grade, s.grade_selector
                        FROM pt_students s
                        JOIN pt_assignments a ON s.id = a.student_id
                        WHERE s.grade = :grade AND s.grade_selector = :grade_selector
                    """),
                    {"grade": grade, "grade_selector": grade_selector}
                ).all()

                # Process Names in Python
                student_data_list = []
                for s in students_query:
                    name_data = names_map.get(s.id, {'first': 'Unknown', 'last': 'Unknown'})
                    student_data_list.append({
                        'id': s.id,
                        'grade': s.grade,
                        'grade_selector': s.grade_selector,
                        'first_name': name_data['first'],
                        'last_name': name_data['last']
                    })

                # Sort by Last Name
                student_data_list.sort(key=lambda x: (x['last_name'], x['first_name']))

                if not student_data_list:
                    continue

                story.append(Paragraph(f"PT-Kurszuordnungen - Klasse {grade}/{grade_selector}", title_style))
                story.append(PageBreak())

                for i, student in enumerate(student_data_list):
                    story.append(Paragraph(f"{student['last_name']}, {student['first_name']}", student_title_style))
                    story.append(Paragraph(f"Klasse: {student['grade']} | ID: {student['grade_selector']}", info_style))
                    story.append(Spacer(1, 0.5*cm))

                    assignments = db.session.execute(
                        text("""
                            SELECT p.title, p.presenter, p.teacher, p.slot, p.room
                            FROM pt_assignments a
                            JOIN pt_presentations p ON a.presentation_id = p.id
                            WHERE a.student_id = :sid
                            ORDER BY p.slot
                        """), {"sid": student['id']}
                    ).all()

                    if assignments:
                        table_data = [['Block', 'Kurstitel', 'Referent', 'Lehrer', 'Raum']]
                        for a in assignments:
                            table_data.append([f"Block {a.slot}", a.title, a.presenter or "-", a.teacher or "-", a.room or "-"])
                        
                        table = Table(table_data, colWidths=[2*cm, 6*cm, 3*cm, 3*cm, 2*cm])
                        table.setStyle(TableStyle([
                            ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                            ('GRID', (0,0), (-1,-1), 1, colors.black),
                            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                        ]))
                        story.append(table)
                    else:
                        story.append(Paragraph("Keine Zuordnungen.", info_style))

                    if i < len(student_data_list) - 1:
                        story.append(PageBreak())

                # Build PDF into RAM buffer
                doc.build(story)
                pdf_size = len(class_pdf_buffer.getvalue())
                print(f"Built PDF for class {grade}/{grade_selector}, size: {pdf_size} bytes")

                # Write RAM buffer into ZIP buffer
                pdf_filename = f"PT_Stundenpläne_{grade}_{grade_selector}.pdf"
                zip_file.writestr(pdf_filename, class_pdf_buffer.getvalue())
                print(f"Added {pdf_filename} to ZIP")

        # Finalize ZIP
        zip_buffer.seek(0)
        zip_size = len(zip_buffer.getvalue())
        print(f"ZIP file created, total size: {zip_size} bytes")

        return send_file(
            zip_buffer,
            as_attachment=True,
            download_name="PT_Schüler_Stundenpläne_Secure.zip",
            mimetype="application/zip"
        )

    except Exception as e:
        error_msg = f"PDF Export Error: {str(e)}"
        print(error_msg)
        import traceback
        traceback.print_exc()
        session['upload_error'] = error_msg
        return redirect("/admin/pt/panel")
    
# Version that requires the Names to be Saved in the DB
@pt_admin_views.route("/pt/export_room_lists_pdf", methods=["GET"])
@admin_required
def pt_export_room_lists_pdf():
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        # Create download directory with absolute path
        current_dir = os.path.dirname(os.path.abspath(__file__))
        download_dir = os.path.join(current_dir, 'data', 'pt', 'downloads')
        os.makedirs(download_dir, exist_ok=True)
        
        print(f"Download directory: {download_dir}")
        
        # Create PDF file path
        pdf_filename = 'PT_Raum_Listen.pdf'
        pdf_path = os.path.join(download_dir, pdf_filename)
        
        print(f"PDF path: {pdf_path}")
        
        # Get all room assignments organized by room and slot
        room_assignments_query = db.session.execute(
            text("""
                SELECT 
                    p.room, p.slot, p.title, p.presenter, p.teacher, p.description,
                    s.first_name, s.last_name, s.grade, s.grade_selector
                FROM pt_assignments a
                JOIN pt_students s ON a.student_id = s.id
                JOIN pt_presentations p ON a.presentation_id = p.id
                ORDER BY p.room, p.slot, s.last_name, s.first_name
            """)
        ).all()
        
        if not room_assignments_query:
            print("No room assignments found")
            # Create a simple PDF with no assignments message
            doc = SimpleDocTemplate(pdf_path, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
            styles = getSampleStyleSheet()
            story = [Paragraph("Keine PT-Raumzuordnungen gefunden.", styles['Title'])]
            doc.build(story)
        else:
            print(f"Found {len(room_assignments_query)} room assignments")
            
            # Organize data by room and slot
            room_data = {}
            for assignment in room_assignments_query:
                room, slot, title, presenter, teacher, description, first_name, last_name, grade, grade_selector = assignment
                
                if room not in room_data:
                    room_data[room] = {}
                
                if slot not in room_data[room]:
                    room_data[room][slot] = {
                        'course_info': {
                            'title': title,
                            'presenter': presenter,
                            'teacher': teacher,
                            'description': description
                        },
                        'students': []
                    }
                
                room_data[room][slot]['students'].append({
                    'first_name': first_name,
                    'last_name': last_name,
                    'grade': grade,
                    'grade_selector': grade_selector 
                })
            
            # Create PDF document
            doc = SimpleDocTemplate(pdf_path, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
            story = []
            
            # Get styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=TA_CENTER,
                textColor=colors.black
            )
            
            room_title_style = ParagraphStyle(
                'RoomTitle',
                parent=styles['Heading2'],
                fontSize=18,
                spaceAfter=20,
                alignment=TA_CENTER,
                textColor=colors.black
            )
            
            course_title_style = ParagraphStyle(
                'CourseTitle',
                parent=styles['Heading3'],
                fontSize=14,
                spaceAfter=15,
                alignment=TA_LEFT,
                textColor=colors.black 
            )
            
            info_style = ParagraphStyle(
                'InfoStyle',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=10,
                alignment=TA_LEFT
            )
            
            # Add main title page
            story.append(Paragraph("PT-Raumzuordnungslisten", title_style))
            story.append(Spacer(1, 2*cm))
            story.append(Paragraph(f"Raumlisten - {len(room_data)} Räume", styles['Heading2']))
            story.append(PageBreak())
            
            # Generate pages for each room and slot combination
            room_count = 0
            total_rooms = sum(len(slots) for slots in room_data.values())
            
            for room, slots in sorted(room_data.items()):
                for slot, slot_data in sorted(slots.items()):
                    room_count += 1
                    
                    # Room and slot header
                    story.append(Paragraph(f"Raum: {room}", room_title_style))
                    story.append(Paragraph(f"Block {slot}", course_title_style))
                    story.append(Spacer(1, 0.5*cm))
                    
                    # Course information
                    course = slot_data['course_info']
                    story.append(Paragraph(f"<b>Kurs:</b> {course['title']}", info_style))
                    story.append(Paragraph(f"<b>Referent:</b> {course['presenter'] or 'N/A'}", info_style))
                    story.append(Paragraph(f"<b>Lehrer:</b> {course['teacher'] or 'N/A'}", info_style))
                    
                    if course['description'] and course['description'].strip():
                        story.append(Paragraph(f"<b>Beschreibung:</b> {course['description']}", info_style))
                    
                    story.append(Spacer(1, 0.5*cm))
                    
                    # Student list
                    students = slot_data['students']
                    story.append(Paragraph(f"<b>Schüler ({len(students)}):</b>", course_title_style))
                    
                    if students:

                        table_data = [
                            ['Nr.', 'Nachname', 'Vorname', 'Klasse', 'Klassen-Bez.', 'Anwesend']
                        ]
                        
                        for i, student in enumerate(students, 1):
                            table_data.append([
                                str(i),
                                student['last_name'] or 'N/A',
                                student['first_name'] or 'N/A',
                                str(student['grade']) or 'N/A',
                                str(student['grade_selector']) or 'N/A',
                                "" 
                            ])
                        
                        # --- ADDED WIDTH FOR NEW COLUMN ---
                        table = Table(table_data, colWidths=[1*cm, 3.5*cm, 3.5*cm, 2*cm, 2.5*cm, 2.5*cm])

                        table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey), 
                            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),    
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, 0), 11),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                            ('FONTSIZE', (0, 1), (-1, -1), 9),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black),
                            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                            ('ALIGN', (0, 0), (0, -1), 'CENTER'), 
                            ('ALIGN', (1, 1), (2, -1), 'LEFT'),
                            ('ALIGN', (3, 1), (-1, -1), 'CENTER'),
                        ]))
                        
                        story.append(table)
                    else:
                        story.append(Paragraph("Keine Schüler diesem Kurs zugeordnet.", info_style))
                    
                    # Add signature line
                    story.append(Spacer(1, 1*cm))
                    story.append(Paragraph("Unterschrift Lehrer: ________________________", info_style))
                    
                    # Add page break except for the last room/slot combination
                    if room_count < total_rooms:
                        story.append(PageBreak())
            
            # Build PDF
            doc.build(story)
        
        # Verify the file was created
        if not os.path.exists(pdf_path):
            print(f"PDF file was not created at {pdf_path}")
            return redirect("/admin/pt/panel")
        
        print(f"PDF created successfully at {pdf_path}")
        print(f"File size: {os.path.getsize(pdf_path)} bytes")

        # Send file using absolute path
        return send_from_directory(
            download_dir,
            pdf_filename,
            as_attachment=True,
            mimetype="application/pdf"
        )
        
    except ImportError as e:
        print(f"Missing reportlab library: {e}")
        print("Please install reportlab: pip install reportlab")
        return redirect("/admin/pt/panel")
    except Exception as e:
        print(f"Error exporting room lists PDF: {e}")
        import traceback
        traceback.print_exc()
        return redirect("/admin/pt/panel")
    
# Version that loads names from uploaded file in RAM ONLY and does not require Names to be saved in DB 
@pt_admin_views.route("/pt/export_room_lists_pdf_secure", methods=["POST"])
@admin_required
def pt_export_room_lists_pdf_secure():
    if "file" not in request.files:
        session['upload_error'] = "No file uploaded for room lists PDF export"
        return redirect("/admin/pt/panel")

    file = request.files["file"]
    if file.filename == "":
        session['upload_error'] = "Empty filename for room lists PDF export"
        return redirect("/admin/pt/panel")

    try:
        # Check for required libraries first
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER
            print("ReportLab imported successfully for room lists")
        except ImportError as ie:
            error_msg = f"Missing required library: {str(ie)}. Please install reportlab: uv pip install reportlab"
            print(f"ERROR: {error_msg}")
            session['upload_error'] = error_msg
            return redirect("/admin/pt/panel")

        from .pt_filehandler import load_names_map
        names_map = load_names_map(file)
        print(f"Loaded {len(names_map)} names into memory for room lists")

        if not names_map:
            session['upload_error'] = "Failed to load names from Excel file. Please check file format (ID, Last Name, First Name in columns A, B, C)"
            return redirect("/admin/pt/panel")

        # Get Data
        raw_data = db.session.execute(text("""
            SELECT p.room, p.slot, p.title, p.presenter, p.teacher, p.description,
                   s.id, s.grade, s.grade_selector
            FROM pt_assignments a
            JOIN pt_students s ON a.student_id = s.id
            JOIN pt_presentations p ON a.presentation_id = p.id
            ORDER BY p.room, p.slot
        """)).all()

        # Organize Data
        room_data = {}
        for row in raw_data:
            room, slot, title, presenter, teacher, description, s_id, grade, g_sel = row
            
            # Name Lookup
            name = names_map.get(s_id, {'first': 'Unknown', 'last': 'Unknown'})
            
            if room not in room_data: room_data[room] = {}
            if slot not in room_data[room]:
                room_data[room][slot] = {
                    'info': {'title': title, 'presenter': presenter, 'teacher': teacher, 'desc': description},
                    'students': []
                }
            
            room_data[room][slot]['students'].append({
                'last': name['last'], 'first': name['first'], 
                'grade': grade, 'sel': g_sel
            })

        # Sort students
        for r in room_data:
            for s in room_data[r]:
                room_data[r][s]['students'].sort(key=lambda x: (x['last'], x['first']))

        # PDF Generation in RAM
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(pdf_buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
        story = []
        styles = getSampleStyleSheet()
        
        # Styles
        title_style = ParagraphStyle('T', parent=styles['Heading1'], fontSize=24, alignment=TA_CENTER, spaceAfter=30)
        room_style = ParagraphStyle('R', parent=styles['Heading2'], fontSize=18, alignment=TA_CENTER, spaceAfter=20)
        course_style = ParagraphStyle('C', parent=styles['Heading3'], fontSize=14, spaceAfter=15)
        text_style = ParagraphStyle('Tx', parent=styles['Normal'], fontSize=12, spaceAfter=10)

        story.append(Paragraph("PT-Raumzuordnungslisten", title_style))
        story.append(PageBreak())

        total_rooms = sum(len(slots) for slots in room_data.values())
        count = 0

        for room, slots in sorted(room_data.items()):
            for slot, data in sorted(slots.items()):
                count += 1
                story.append(Paragraph(f"Raum: {room}", room_style))
                story.append(Paragraph(f"Block {slot} - {data['info']['title']}", course_style))
                story.append(Paragraph(f"Lehrer: {data['info']['teacher'] or '-'}", text_style))
                story.append(Spacer(1, 0.5*cm))

                students = data['students']
                if students:
                    table_data = [['Nr.', 'Nachname', 'Vorname', 'Klasse', 'Bez.', 'Check']]
                    for i, st in enumerate(students, 1):
                        table_data.append([str(i), st['last'], st['first'], str(st['grade']), str(st['sel'] or ''), ''])
                    
                    t = Table(table_data, colWidths=[1*cm, 4*cm, 4*cm, 2*cm, 2*cm, 2*cm])
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                        ('GRID', (0,0), (-1,-1), 1, colors.black),
                        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                        ('ALIGN', (1,1), (2,-1), 'LEFT'), # Left align names
                    ]))
                    story.append(t)
                else:
                    story.append(Paragraph("Keine Schüler.", text_style))

                story.append(Spacer(1, 1*cm))
                story.append(Paragraph("Unterschrift: __________________", text_style))
                
                if count < total_rooms:
                    story.append(PageBreak())

        doc.build(story)
        pdf_buffer.seek(0)

        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name="PT_Raum_Listen_Secure.pdf",
            mimetype="application/pdf"
        )

    except Exception as e:
        error_msg = f"Room Lists PDF Export Error: {str(e)}"
        print(error_msg)
        import traceback
        traceback.print_exc()
        session['upload_error'] = error_msg
        return redirect("/admin/pt/panel")