#subdomain\src\app\pt_logincode_export.py
from docx import Document
from docx.shared import Pt
from sqlalchemy import text
import zipfile
import os

from . import db


def get_class_with_names(grade, grade_selector):
    # Get student data directly from pt_students table with names
    query = db.session.execute(
        text(
            f"SELECT first_name, last_name, logincode FROM pt_students WHERE grade = {grade} AND grade_selector = {grade_selector} ORDER BY last_name, first_name"
        )
    ).fetchall()
    return query


def get_sek2_with_names(grade):
    # Get student data directly from pt_students table with names
    query = db.session.execute(
        text(
            f"SELECT first_name, last_name, logincode FROM pt_students WHERE grade = {grade} ORDER BY last_name, first_name"
        )
    ).fetchall()
    return query


def zip_files():
    directory = os.getenv("OUTPUT_DIR", "./app/data/pt/downloads")
    zip_file_path = os.path.join(directory, "PT_Logincodes.zip")

    if not os.path.exists(directory):
        print(f"Directory {directory} does not exist.")
        return

    with zipfile.ZipFile(zip_file_path, "w") as zipf:
        for file in os.listdir(directory):
            if file.endswith(".docx"):
                zipf.write(os.path.join(directory, file), file)

    print("Zipped files successfully.")


def export_logincodes():
    print(os.getcwd())
    output_dir = os.getenv("OUTPUT_DIR", "./app/data/pt/downloads")
    os.makedirs(output_dir, exist_ok=True)

    print(f"Cleaning old files from {output_dir}...")
    for file in os.listdir(output_dir):
        if file.endswith(".docx") or file == "PT_Logincodes.zip":
            try:
                os.remove(os.path.join(output_dir, file))
            except Exception as e:
                print(f"Could not delete file {file}: {e}")
                
    for grade in range(5, 13, 1):
        
        # -----------------------------------------------------------
        # CASE 1: Grades exported as a whole (5, 6, 11, 12)
        # -----------------------------------------------------------
        if grade in [5, 6, 11, 12]:
            print(f"Exporting grade {grade}")
            query = get_sek2_with_names(grade)
            
            if len(query) > 0:
                doc = Document()
                doc.add_heading(f"PT Login Codes {grade}", 0)
                
                # Table setup
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                
                headers = ["First Name", "Last Name", "Login Code"]
                for i, header in enumerate(headers):
                    cell = table.cell(0, i)
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        run = paragraph.runs[0]
                        run.bold = True
                        run.font.size = Pt(14)

                for row in query:
                    row_cells = table.add_row().cells
                    row_cells[0].text = row.first_name
                    row_cells[1].text = row.last_name
                    row_cells[2].text = row.logincode

                for row in table.rows[1:]:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.space_after = Pt(14)
                            paragraph.paragraph_format.space_before = Pt(14)
                            paragraph.paragraph_format.line_spacing = Pt(14)

                doc.save(os.path.join(output_dir, f"PT_Logincodes_{grade}.docx"))
                print(f"Finished grade {grade}")
            
            # Skip the rest of the loop for these grades so we don't look for selectors
            continue 

        # -----------------------------------------------------------
        # CASE 2: Grades split by Selector (7, 8, 9, 10, etc.)
        # -----------------------------------------------------------
        
        # dynamic query: find out exactly which selectors exist for this grade
        existing_selectors_query = db.session.execute(
            text(f"SELECT DISTINCT grade_selector FROM pt_students WHERE grade = {grade} ORDER BY grade_selector ASC")
        ).fetchall()
        
        # Convert to a simple list, filtering out None
        active_selectors = [r[0] for r in existing_selectors_query if r[0] is not None]

        for grade_selector in active_selectors:
            print(f"Exporting grade {grade}/{grade_selector}")
            query = get_class_with_names(grade, grade_selector)

            if len(query) > 0:
                doc = Document()
                doc.add_heading(f"PT Login Codes {grade}/{grade_selector}", 0)

                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                
                headers = ["First Name", "Last Name", "Login Code"]
                for i, header in enumerate(headers):
                    cell = table.cell(0, i)
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        run = paragraph.runs[0]
                        run.bold = True
                        run.font.size = Pt(14)

                for row in query:
                    row_cells = table.add_row().cells
                    row_cells[0].text = row.first_name
                    row_cells[1].text = row.last_name
                    row_cells[2].text = row.logincode

                for row in table.rows[1:]:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.space_after = Pt(14)
                            paragraph.paragraph_format.space_before = Pt(14)
                            paragraph.paragraph_format.line_spacing = Pt(14)

                doc.save(os.path.join(output_dir, f"PT_Logincodes_{grade}_{grade_selector}.docx"))
                print(f"Finished grade {grade}/{grade_selector}")

    zip_files()