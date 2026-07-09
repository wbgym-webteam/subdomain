#subdomain\src\app\tdw_logincode_export_secure.py
# This secure version reads names from the Excel file in RAM during export
# Names are never stored in the database, only matched by ID
# Files are generated entirely in RAM and never saved to disk

from docx import Document
from docx.shared import Pt
from sqlalchemy import text
import zipfile
import io

from . import db


def export_logincodes_secure_ram(names_map):
    """
    Export login codes using names from RAM only.
    Generates all files in RAM and returns a BytesIO buffer containing the ZIP.
    names_map: dictionary with student_id as key and {'first': ..., 'last': ...} as value
    """
    # Create ZIP buffer in RAM
    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:

        for grade in range(5, 13, 1):

            # ----------------------------------------------------------------
            # CASE 1: Grades exported as a whole (5, 6, 11, 12)
            if grade in [5, 6]:
                print(f"Exporting grade {grade}")

                # Get student IDs and login codes from DB
                students = db.session.execute(
                    text(f"SELECT id, logincode FROM students WHERE grade = {grade}")
                ).fetchall()

                if not students:
                    continue

                # Match with names from RAM and sort
                student_data = []
                for s in students:
                    if s.id in names_map:
                        student_data.append({
                            'first': names_map[s.id]['first'],
                            'last': names_map[s.id]['last'],
                            'code': s.logincode
                        })

                # Sort by last name, then first name
                student_data.sort(key=lambda x: (x['last'], x['first']))

                if student_data:
                    # Generate DOCX in RAM
                    doc = Document()
                    doc.add_heading(f"TDW Login Codes {grade}", 0)

                    table = doc.add_table(rows=1, cols=4)
                    table.style = "Table Grid"

                    # Headers
                    headers = ["Website", "First Name", "Last Name", "Login Code"]
                    for i, h in enumerate(headers):
                        cell = table.cell(0, i)
                        cell.text = h
                        cell.paragraphs[0].runs[0].bold = True
                        cell.paragraphs[0].runs[0].font.size = Pt(14)

                    # Rows
                    for s in student_data:
                        row_cells = table.add_row().cells
                        row_cells[0].text = "tdw.wbgym.de"
                        row_cells[1].text = s['first']
                        row_cells[2].text = s['last']
                        row_cells[3].text = s['code']

                        # Apply spacing
                        for cell in row_cells:
                            for p in cell.paragraphs:
                                p.paragraph_format.space_after = Pt(14)
                                p.paragraph_format.space_after = Pt(14)
                                p.paragraph_format.space_before = Pt(14)
                                p.paragraph_format.line_spacing = Pt(14)

                    # Save DOCX to RAM buffer
                    docx_buffer = io.BytesIO()
                    doc.save(docx_buffer)
                    zip_file.writestr(f"TdW_Logincodes_{grade}.docx", docx_buffer.getvalue())
                    print(f"Finished grade {grade}")

                continue  # Move to next grade

            # ----------------------------------------------------------------
            # CASE 2: Grades split by selector (7, 8, 9, 10, etc.)
            existing_selectors_query = db.session.execute(
                text(f"SELECT DISTINCT grade_selector FROM students WHERE grade = {grade} ORDER BY grade_selector ASC")
            ).fetchall()

            active_selectors = [r[0] for r in existing_selectors_query if r[0] is not None]

            for grade_selector in active_selectors:
                print(f"Exporting grade {grade}/{grade_selector}")

                # Get student IDs and login codes from DB
                students = db.session.execute(
                    text(f"SELECT id, logincode FROM students WHERE grade = {grade} AND grade_selector = {grade_selector}")
                ).fetchall()

                if not students:
                    continue

                # Match with names from RAM and sort
                student_data = []
                for s in students:
                    if s.id in names_map:
                        student_data.append({
                            'first': names_map[s.id]['first'],
                            'last': names_map[s.id]['last'],
                            'code': s.logincode
                        })

                # Sort by last name, then first name
                student_data.sort(key=lambda x: (x['last'], x['first']))

                if student_data:
                    # Generate DOCX in RAM
                    doc = Document()
                    doc.add_heading(f"TDW Login Codes {grade}/{grade_selector}", 0)

                    table = doc.add_table(rows=1, cols=4)
                    table.style = "Table Grid"

                    # Headers
                    headers = ["Website", "First Name", "Last Name", "Login Code"]
                    for i, h in enumerate(headers):
                        cell = table.cell(0, i)
                        cell.text = h
                        cell.paragraphs[0].runs[0].bold = True
                        cell.paragraphs[0].runs[0].font.size = Pt(14)

                    # Rows
                    for s in student_data:
                        row_cells = table.add_row().cells
                        row_cells[0].text = "tdw.wbgym.de"
                        row_cells[1].text = s['first']
                        row_cells[2].text = s['last']
                        row_cells[3].text = s['code']

                        # Apply spacing
                        for cell in row_cells:
                            for p in cell.paragraphs:
                                p.paragraph_format.space_after = Pt(14)
                                p.paragraph_format.space_before = Pt(14)
                                p.paragraph_format.line_spacing = Pt(14)

                    # Save DOCX to RAM buffer
                    docx_buffer = io.BytesIO()
                    doc.save(docx_buffer)
                    zip_file.writestr(f"TdW_Logincodes_{grade}_{grade_selector}.docx", docx_buffer.getvalue())
                    print(f"Finished grade {grade}/{grade_selector}")

    # Finalize ZIP and return buffer
    zip_buffer.seek(0)
    return zip_buffer
