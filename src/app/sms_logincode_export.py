from docx import Document
from docx.shared import Pt
from sqlalchemy import text
from openpyxl import load_workbook
import zipfile
import io
from collections import defaultdict

from . import db


def load_names_map(file_storage):
    workbook = load_workbook(file_storage, keep_vba=False, data_only=True)
    sheet = workbook.worksheets[0]
    names_map = {}
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if row[0] is not None:
            names_map[row[0]] = {
                'first': row[2] or "",
                'last': row[1] or "",
            }
    return names_map


def _load_all_students(names_map):
    """Load all students in one query, grouped by (grade, grade_selector)."""
    rows = db.session.execute(
        text("SELECT Student_id, logincode, grade, grade_selector FROM students_sms")
    ).fetchall()

    groups = defaultdict(list)
    for student_id, logincode, grade, grade_selector in rows:
        n = names_map.get(student_id, {'first': '', 'last': ''})
        groups[(grade, grade_selector)].append((n['last'], n['first'], logincode))

    # Sort each group by last name, first name
    for key in groups:
        groups[key].sort()

    return groups


def export_logincodes(names_file):
    names_map = load_names_map(names_file)
    groups = _load_all_students(names_map)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        for grade in range(5, 13):
            special_grade = grade in (5, 6, 12)

            if special_grade:
                # Merge all grade_selectors for this grade
                students = []
                for (g, gs), recs in groups.items():
                    if g == grade:
                        students.extend(recs)
                students.sort()
                if not students:
                    continue
                doc = Document()
                doc.add_heading(f"SMS Login Codes {grade}", 0)
                _fill_table(doc, students)
                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                zipf.writestr(f"SmS_Logincodes_{grade}.docx", docx_buffer.getvalue())
            else:
                selectors = sorted(gs for (g, gs) in groups if g == grade)
                for grade_selector in selectors:
                    students = groups.get((grade, grade_selector), [])
                    if not students:
                        continue
                    doc = Document()
                    doc.add_heading(f"SMS Login Codes {grade}/{grade_selector}", 0)
                    _fill_table(doc, students)
                    docx_buffer = io.BytesIO()
                    doc.save(docx_buffer)
                    zipf.writestr(f"SmS_Logincodes_{grade}_{grade_selector}.docx", docx_buffer.getvalue())

    zip_buffer.seek(0)
    return zip_buffer


def _fill_table(doc, students):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["First Name", "Last Name", "Login Code", "Link"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(14)

    space = Pt(14)
    for last_name, first_name, logincode in students:
        row_cells = table.add_row().cells
        row_cells[0].text = first_name
        row_cells[1].text = last_name
        row_cells[2].text = logincode
        row_cells[3].text = "sms.wbgym.de"
        for cell in row_cells:
            fmt = cell.paragraphs[0].paragraph_format
            fmt.space_after = space
            fmt.space_before = space
            fmt.line_spacing = space
